import json

import docx
import numpy as np

from resume_ranker.app import pipeline as pipeline_module
from resume_ranker.app.pipeline import RankingPipeline
from resume_ranker.config import config


class FakeEmbedder:
    embedding_mode = "fake-local"
    using_fallback = False

    def __init__(self, app_config):
        self.config = app_config

    def fit_reference_corpus(self, reference_texts):
        return None

    def encode(self, texts, batch_size):
        vectors = []
        for text in texts:
            lower = text.lower()
            vectors.append([
                1.0 if "python" in lower or "faiss" in lower else 0.2,
                1.0 if "engineer" in lower else 0.1,
            ])
        arr = np.asarray(vectors, dtype=np.float32)
        norms = np.linalg.norm(arr, axis=1, keepdims=True)
        norms[norms == 0] = 1.0
        return arr / norms

    def encode_single(self, text):
        return self.encode([text], batch_size=1)[0]


def _write_docx(path):
    document = docx.Document()
    document.add_heading("Senior AI Engineer", level=1)
    document.add_paragraph("We need Python, FAISS, retrieval, and ranking experience.")
    document.save(str(path))


def _write_candidates(path, candidate_dict_factory):
    rows = [
        candidate_dict_factory(candidate_id="CAND_0000001"),
        candidate_dict_factory(
            candidate_id="CAND_0000002",
            profile={"headline": "Backend Engineer", "current_title": "Backend Engineer"},
            skills=[{"name": "Java", "proficiency": "advanced", "endorsements": 2, "duration_months": 24}],
        ),
        candidate_dict_factory(
            candidate_id="CAND_0000003",
            skills=[{"name": "Rust", "proficiency": "expert", "endorsements": 0, "duration_months": 0}],
        ),
    ]
    with open(path, "w", encoding="utf-8") as f:
        for row in rows:
            f.write(json.dumps(row) + "\n")


def test_pipeline_writes_submission_and_audit(tmp_path, candidate_dict_factory, monkeypatch):
    monkeypatch.setattr(pipeline_module, "LocalEmbedder", FakeEmbedder)
    candidates_path = tmp_path / "candidates.jsonl"
    jd_path = tmp_path / "job_description.docx"
    output_path = tmp_path / "submission.csv"
    _write_candidates(candidates_path, candidate_dict_factory)
    _write_docx(jd_path)

    app_config = config.model_copy(
        update={
            "top_n": 2,
            "faiss_top_k": 2,
            "artifacts_dir": tmp_path / "artifacts",
        }
    )

    df = RankingPipeline(app_config).run(candidates_path, jd_path, output_path)

    assert output_path.exists()
    assert output_path.with_suffix(".audit.json").exists()
    assert list(df.columns) == app_config.csv_columns
    assert len(df) == 2
    assert df["candidate_id"].to_list()[0] == "CAND_0000001"

    audit = json.loads(output_path.with_suffix(".audit.json").read_text(encoding="utf-8"))
    assert audit["candidate_counts"]["total"] == 3
    assert audit["candidate_counts"]["honeypots_excluded"] == 1
    assert audit["validation"]["passed"] is True
