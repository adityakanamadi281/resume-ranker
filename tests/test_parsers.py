"""tests/test_parsers.py -- unit tests for docx_parser and candidate_parser."""

import json
import logging
from pathlib import Path

import docx
import pytest
from pydantic import ValidationError

from src.parsers.candidate_parser import stream_candidates
from src.parsers.docx_parser import parse_job_description
from src.schema import Candidate


# -----------------------------------------------------------------------
# DOCX parser
# -----------------------------------------------------------------------

def _write_mock_docx(path: Path) -> None:
    document = docx.Document()
    document.add_heading("Job Description: Test Role", level=0)
    document.add_paragraph("Section One Header")
    document.add_paragraph("This is the first paragraph of content.")
    document.add_paragraph("Section Two Header")
    document.add_paragraph("This is the second paragraph of content.")
    document.save(str(path))


def test_docx_parser_extracts_all_paragraphs(tmp_path):
    docx_path = tmp_path / "jd.docx"
    _write_mock_docx(docx_path)

    text = parse_job_description(docx_path)

    assert "This is the first paragraph of content." in text
    assert "This is the second paragraph of content." in text


def test_docx_parser_preserves_section_headers(tmp_path):
    docx_path = tmp_path / "jd.docx"
    _write_mock_docx(docx_path)

    text = parse_job_description(docx_path)

    assert "Section One Header" in text
    assert "Section Two Header" in text
    assert text.index("Section One Header") < text.index("This is the first paragraph")
    assert text.index("Section Two Header") < text.index("This is the second paragraph")


def test_docx_parser_missing_file_raises(tmp_path):
    missing_path = tmp_path / "does_not_exist.docx"
    with pytest.raises(FileNotFoundError):
        parse_job_description(missing_path)


def test_docx_parser_invalid_file_raises(tmp_path):
    not_a_docx = tmp_path / "not_a_docx.docx"
    not_a_docx.write_text("this is plain text, not a real docx file")
    with pytest.raises(ValueError):
        parse_job_description(not_a_docx)


# -----------------------------------------------------------------------
# Candidate parser (streaming)
# -----------------------------------------------------------------------

def _write_jsonl(path: Path, lines):
    with open(path, "w", encoding="utf-8") as f:
        for line in lines:
            f.write(line + "\n")


def test_stream_candidates_yields_correct_objects(tmp_path, candidate_dict_factory):
    jsonl_path = tmp_path / "candidates.jsonl"
    valid_dict = candidate_dict_factory()
    _write_jsonl(jsonl_path, [json.dumps(valid_dict)])

    results = list(stream_candidates(jsonl_path))

    assert len(results) == 1
    assert isinstance(results[0], Candidate)
    assert results[0].candidate_id == valid_dict["candidate_id"]


def test_stream_candidates_skips_invalid_schema(tmp_path, candidate_dict_factory, caplog):
    jsonl_path = tmp_path / "candidates.jsonl"
    valid_dict = candidate_dict_factory()
    invalid_dict = candidate_dict_factory(candidate_id="NOT_VALID_ID")
    _write_jsonl(jsonl_path, [json.dumps(valid_dict), json.dumps(invalid_dict)])

    with caplog.at_level(logging.WARNING):
        results = list(stream_candidates(jsonl_path))

    assert len(results) == 1
    assert results[0].candidate_id == valid_dict["candidate_id"]
    assert any("failed schema validation" in r.message for r in caplog.records)


def test_stream_candidates_handles_malformed_json(tmp_path, candidate_dict_factory, caplog):
    jsonl_path = tmp_path / "candidates.jsonl"
    valid_dict = candidate_dict_factory()
    _write_jsonl(jsonl_path, [json.dumps(valid_dict), "{not valid json", ""])

    with caplog.at_level(logging.WARNING):
        results = list(stream_candidates(jsonl_path))

    assert len(results) == 1
    assert any("invalid JSON" in r.message for r in caplog.records)


def test_stream_candidates_missing_file_raises(tmp_path):
    with pytest.raises(FileNotFoundError):
        list(stream_candidates(tmp_path / "missing.jsonl"))


# -----------------------------------------------------------------------
# Schema validation
# -----------------------------------------------------------------------

def test_valid_candidate_passes_validation(candidate_dict_factory):
    candidate = Candidate.model_validate(candidate_dict_factory())
    assert candidate.candidate_id == "CAND_0000001"


def test_invalid_candidate_id_format_fails(candidate_dict_factory):
    with pytest.raises(ValidationError):
        Candidate.model_validate(candidate_dict_factory(candidate_id="not-a-valid-id"))


def test_negative_notice_period_fails(candidate_dict_factory):
    raw = candidate_dict_factory(redrob_signals={"notice_period_days": -5})
    with pytest.raises(ValidationError):
        Candidate.model_validate(raw)


def test_out_of_range_response_rate_fails(candidate_dict_factory):
    raw = candidate_dict_factory(redrob_signals={"recruiter_response_rate": 1.5})
    with pytest.raises(ValidationError):
        Candidate.model_validate(raw)
