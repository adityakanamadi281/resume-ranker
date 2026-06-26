import json
from typer.testing import CliRunner
import docx
import numpy as np

from resume_ranker.cli.rank import app as rank_app
from resume_ranker.cli.precompute import app as precompute_app
from resume_ranker.cli.validate import app as validate_app
from resume_ranker.cli.benchmark import app as benchmark_app
from resume_ranker.app import pipeline as pipeline_module


class FakeEmbedder:
    embedding_mode = "fake-local"
    using_fallback = False

    def __init__(self, app_config):
        self.config = app_config

    def encode(self, texts, batch_size):
        vectors = []
        for text in texts:
            lower = text.lower()
            vectors.append([
                1.0 if "python" in lower or "faiss" in lower else 0.2,
                1.0 if "engineer" in lower else 0.1,
            ])
        arr = np.asarray(vectors, dtype=np.float32)
        norms = np.linalg.norm(arr, axis=1, keepdims=True)
        norms[norms == 0] = 1.0
        return arr / norms

    def encode_single(self, text):
        return self.encode([text], batch_size=1)[0]


def _write_docx(path):
    document = docx.Document()
    document.add_heading("Senior AI Engineer", level=1)
    document.add_paragraph("We need Python, FAISS, retrieval, and ranking experience.")
    document.save(str(path))


def _write_candidates(path, candidate_dict_factory):
    rows = [
        candidate_dict_factory(candidate_id="CAND_0000001"),
        candidate_dict_factory(
            candidate_id="CAND_0000002",
            profile={"headline": "Backend Engineer", "current_title": "Backend Engineer"},
            skills=[{"name": "Java", "proficiency": "advanced", "endorsements": 2, "duration_months": 24}],
        ),
    ]
    with open(path, "w", encoding="utf-8") as f:
        for row in rows:
            f.write(json.dumps(row) + "\n")


def test_rank_cli_help():
    runner = CliRunner()
    result = runner.invoke(rank_app, ["--help"])
    assert result.exit_code == 0
    assert "Rank candidates against a job description" in result.stdout


def test_rank_cli_missing_candidates(tmp_path):
    runner = CliRunner()
    jd_path = tmp_path / "jd.docx"
    _write_docx(jd_path)
    result = runner.invoke(
        rank_app,
        ["--candidates", str(tmp_path / "missing.jsonl"), "--jd", str(jd_path)],
    )
    assert result.exit_code == 1
    # Check stderr specifically or stdout
    output = result.stdout or result.stderr
    assert "ERROR: Candidates file not found" in output


def test_rank_cli_missing_jd(tmp_path, candidate_dict_factory):
    runner = CliRunner()
    candidates_path = tmp_path / "candidates.jsonl"
    _write_candidates(candidates_path, candidate_dict_factory)
    result = runner.invoke(
        rank_app,
        ["--candidates", str(candidates_path), "--jd", str(tmp_path / "missing.docx")],
    )
    assert result.exit_code == 1
    output = result.stdout or result.stderr
    assert "ERROR: JD file not found" in output


def test_rank_cli_success(tmp_path, candidate_dict_factory, monkeypatch):
    monkeypatch.setattr(pipeline_module, "LocalEmbedder", FakeEmbedder)
    runner = CliRunner()
    candidates_path = tmp_path / "candidates.jsonl"
    jd_path = tmp_path / "jd.docx"
    output_path = tmp_path / "submission.csv"
    _write_candidates(candidates_path, candidate_dict_factory)
    _write_docx(jd_path)

    # Override config artifacts_dir
    from resume_ranker.cli.rank import config as rank_config
    monkeypatch.setattr(rank_config, "artifacts_dir", tmp_path / "artifacts")
    from resume_ranker.cli.precompute import config as precompute_config
    monkeypatch.setattr(precompute_config, "artifacts_dir", tmp_path / "artifacts")

    # Run precompute first
    precompute_res = runner.invoke(
        precompute_app,
        [
            "--candidates",
            str(candidates_path),
            "--jd",
            str(jd_path),
        ],
    )
    assert precompute_res.exit_code == 0

    # Run rank next
    result = runner.invoke(
        rank_app,
        [
            "--candidates",
            str(candidates_path),
            "--jd",
            str(jd_path),
            "--output",
            str(output_path),
        ],
    )
    assert result.exit_code == 0
    assert "Success! Ranked" in result.stdout
    assert output_path.exists()
    assert output_path.with_suffix(".audit.json").exists()


def test_rank_cli_with_schema_override(tmp_path, candidate_dict_factory, monkeypatch):
    monkeypatch.setattr(pipeline_module, "LocalEmbedder", FakeEmbedder)
    runner = CliRunner()
    candidates_path = tmp_path / "candidates.jsonl"
    jd_path = tmp_path / "jd.docx"
    output_path = tmp_path / "submission.csv"
    schema_path = tmp_path / "custom_schema.json"

    # Override config artifacts_dir
    from resume_ranker.cli.rank import config as rank_config
    monkeypatch.setattr(rank_config, "artifacts_dir", tmp_path / "artifacts")
    from resume_ranker.cli.precompute import config as precompute_config
    monkeypatch.setattr(precompute_config, "artifacts_dir", tmp_path / "artifacts")

    # Write a simple JSON schema that makes CAND_0000002 invalid, but CAND_0000001 valid
    schema = {
        "$schema": "https://json-schema.org/draft/2020-12/schema",
        "type": "object",
        "properties": {
            "candidate_id": {
                "type": "string",
                "pattern": "^CAND_0000001$",  # only CAND_0000001 is allowed
            }
        },
        "required": ["candidate_id"],
    }
    schema_path.write_text(json.dumps(schema))

    _write_candidates(candidates_path, candidate_dict_factory)
    _write_docx(jd_path)

    # Run precompute first
    precompute_res = runner.invoke(
        precompute_app,
        [
            "--candidates",
            str(candidates_path),
            "--jd",
            str(jd_path),
            "--schema",
            str(schema_path),
        ],
    )
    assert precompute_res.exit_code == 0

    # Run rank next
    result = runner.invoke(
        rank_app,
        [
            "--candidates",
            str(candidates_path),
            "--jd",
            str(jd_path),
            "--output",
            str(output_path),
            "--schema",
            str(schema_path),
        ],
    )
    assert result.exit_code == 0
    # The output should only contain CAND_0000001 as CAND_0000002 failed custom schema validation
    # Let's verify by loading the audit log to check candidates counts
    audit_path = output_path.with_suffix(".audit.json")
    assert audit_path.exists()
    audit = json.loads(audit_path.read_text(encoding="utf-8"))
    assert audit["candidate_counts"]["valid_after_honeypot_filter"] == 1


def test_precompute_cli_success(tmp_path, candidate_dict_factory, monkeypatch):
    monkeypatch.setattr(pipeline_module, "LocalEmbedder", FakeEmbedder)
    runner = CliRunner()
    candidates_path = tmp_path / "candidates.jsonl"
    jd_path = tmp_path / "jd.docx"
    _write_candidates(candidates_path, candidate_dict_factory)
    _write_docx(jd_path)

    # Precompute outputs to config.artifacts_dir, let's override that dynamically using monkeypatch on config
    from resume_ranker.cli.precompute import config as precompute_config
    monkeypatch.setattr(precompute_config, "artifacts_dir", tmp_path / "artifacts")

    result = runner.invoke(
        precompute_app,
        [
            "--candidates",
            str(candidates_path),
            "--jd",
            str(jd_path),
        ],
    )
    assert result.exit_code == 0
    assert "Success! Precomputed" in result.stdout or "Precomputation completed" in result.stdout
    assert (tmp_path / "artifacts" / "embeddings.npy").exists()
    assert (tmp_path / "artifacts" / "candidate_ids.npy").exists()


def test_validate_cli(tmp_path):
    runner = CliRunner()
    submission_path = tmp_path / "submission.csv"

    # 1. Missing file should fail validation
    result = runner.invoke(validate_app, ["--submission", str(submission_path)])
    assert result.exit_code == 1
    output = result.stdout or result.stderr
    assert "FAIL" in output

    # 2. Correctly formatted CSV should pass
    submission_path.write_text("candidate_id,rank,score,reasoning\nCAND_0000001,1,0.95,Valid Candidate\n")
    result = runner.invoke(validate_app, ["--submission", str(submission_path)])
    assert result.exit_code == 0
    assert "PASS" in result.stdout


def test_benchmark_cli_help():
    runner = CliRunner()
    result = runner.invoke(benchmark_app, ["--help"])
    assert result.exit_code == 0
    assert "Benchmark the ranking pipeline stages" in result.stdout


def test_benchmark_cli_success(tmp_path, candidate_dict_factory, monkeypatch):
    monkeypatch.setattr(pipeline_module, "LocalEmbedder", FakeEmbedder)
    runner = CliRunner()
    candidates_path = tmp_path / "candidates.jsonl"
    jd_path = tmp_path / "jd.docx"
    output_path = tmp_path / "benchmark.csv"
    _write_candidates(candidates_path, candidate_dict_factory)
    _write_docx(jd_path)

    # Override config artifacts_dir for all clis to stay isolated
    from resume_ranker.cli.benchmark import config as benchmark_config
    monkeypatch.setattr(benchmark_config, "artifacts_dir", tmp_path / "artifacts")
    from resume_ranker.cli.precompute import config as precompute_config
    monkeypatch.setattr(precompute_config, "artifacts_dir", tmp_path / "artifacts")

    # Run precompute first to generate required artifacts
    precompute_res = runner.invoke(
        precompute_app,
        [
            "--candidates",
            str(candidates_path),
            "--jd",
            str(jd_path),
        ],
    )
    assert precompute_res.exit_code == 0

    # Run benchmark command
    result = runner.invoke(
        benchmark_app,
        [
            "--jd",
            str(jd_path),
            "--output",
            str(output_path),
        ],
    )
    assert result.exit_code == 0
    stdout = result.stdout
    assert "Load artifacts" in stdout
    assert "Embed JD" in stdout
    assert "FAISS search" in stdout
    assert "Feature scoring" in stdout
    assert "Explanation" in stdout
    assert "CSV write" in stdout
    assert "Total" in stdout
